"""
RFP Analyzer API — document intelligence endpoints.

POST   /api/rfp-analyzer              Upload document, start async pipeline
GET    /api/rfp-analyzer/history      List user's past analyses
GET    /api/rfp-analyzer/{id}         Poll status / get full report
PATCH  /api/rfp-analyzer/{id}/scope   Override classification for one requirement
DELETE /api/rfp-analyzer/{id}         Hard delete analysis + all child rows
POST   /api/rfp-analyzer/{id}/export/pdf   Export report as PDF
POST   /api/rfp-analyzer/{id}/export/docx  Export report as DOCX
"""
import asyncio
import io
import uuid
from datetime import datetime, timezone

import structlog
from anthropic import AsyncAnthropic
from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select, delete as sa_delete
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.core.utils.query_sanitizer import sanitize_context
from app.core.utils.rate_limiter import limiter
from app.db.session import get_db, AsyncSessionLocal
from app.dependencies import get_anthropic_client, get_current_user
from app.models.db_models import RFPAnalysis, RFPRequirement, RFPClassification, User
from app.core.llm.rfp_analyzer_service import run_rfp_analyzer_pipeline
from app.services.rfp_analyzer_parser import parse_rfp_document, SUPPORTED_TYPES

router = APIRouter()
logger = structlog.get_logger()

MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024  # 20 MB


# ── Request / Response models ─────────────────────────────────────────────────

class StartAnalysisResponse(BaseModel):
    analysis_id: str
    status: str


class ClassificationOut(BaseModel):
    scope: str
    justification: str | None
    confidence: float | None
    conditions: str | None
    user_override: str | None
    override_reason: str | None


class RequirementOut(BaseModel):
    req_id: str
    text: str
    raw_quote: str | None
    category: str | None
    priority: str | None
    source_page: int | None
    source_section: str | None
    classification: ClassificationOut | None


class AnalysisReportOut(BaseModel):
    analysis_id: str
    status: str
    original_name: str
    client_name: str | None
    country: str | None
    sector: str | None
    tender_id: str | None
    submission_deadline: str | None
    evaluation_split: dict | None
    budget_indication: str | None
    currency: str | None
    language: str | None
    error_message: str | None
    requirements: list[RequirementOut]
    created_at: str


class AnalysisListItem(BaseModel):
    analysis_id: str
    status: str
    original_name: str
    client_name: str | None
    created_at: str
    requirement_count: int


class OverrideRequest(BaseModel):
    req_id: str          # The req_id string (e.g. "REQ-001"), not the DB UUID
    scope: str           # Must be 'in', 'conditional', or 'out'
    reason: str | None = None


# ── Helpers ───────────────────────────────────────────────────────────────────

def _analysis_to_report(analysis: RFPAnalysis) -> AnalysisReportOut:
    reqs = []
    for req in analysis.requirements:
        c = req.classification
        reqs.append(RequirementOut(
            req_id=req.req_id,
            text=req.text,
            raw_quote=req.raw_quote,
            category=req.category,
            priority=req.priority,
            source_page=req.source_page,
            source_section=req.source_section,
            classification=ClassificationOut(
                scope=c.user_override or c.scope,
                justification=c.justification,
                confidence=c.confidence,
                conditions=c.conditions,
                user_override=c.user_override,
                override_reason=c.override_reason,
            ) if c else None,
        ))

    return AnalysisReportOut(
        analysis_id=str(analysis.id),
        status=analysis.status,
        original_name=analysis.original_name,
        client_name=analysis.client_name,
        country=analysis.country,
        sector=analysis.sector,
        tender_id=analysis.tender_id,
        submission_deadline=analysis.submission_deadline,
        evaluation_split=analysis.evaluation_split,
        budget_indication=analysis.budget_indication,
        currency=analysis.currency,
        language=analysis.language,
        error_message=analysis.error_message,
        requirements=reqs,
        created_at=analysis.created_at.isoformat(),
    )


async def _get_owned_analysis(
    analysis_id: str,
    current_user: User,
    db: AsyncSession,
) -> RFPAnalysis:
    """Fetch analysis with IDOR protection — user must own the row."""
    try:
        aid = uuid.UUID(analysis_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid analysis_id format")

    result = await db.execute(
        select(RFPAnalysis)
        .where(RFPAnalysis.id == aid, RFPAnalysis.user_id == current_user.id)
        .options(selectinload(RFPAnalysis.requirements).selectinload(RFPRequirement.classification))
    )
    analysis = result.scalar_one_or_none()
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return analysis


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("", response_model=StartAnalysisResponse)
@limiter.limit("5/minute")
async def start_analysis(
    request: Request,
    file: UploadFile = File(...),
    analysis_id: str = Form(...),         # Client-generated UUID (spec rule 9.8)
    company_context: str | None = Form(default=None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
    anthropic_client: AsyncAnthropic = Depends(get_anthropic_client),
):
    # Validate analysis_id is a UUID
    try:
        aid = uuid.UUID(analysis_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="analysis_id must be a valid UUID v4")

    # Validate file type
    original_name = file.filename or "unknown"
    ext = original_name.rsplit(".", 1)[-1].lower() if "." in original_name else ""
    if ext not in SUPPORTED_TYPES:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file type '.{ext}'. Supported: {sorted(SUPPORTED_TYPES)}",
        )

    # Read and validate file size
    file_bytes = await file.read()
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File exceeds 20 MB limit ({len(file_bytes) // 1024 // 1024} MB received)",
        )

    # Sanitize company_context
    clean_context = sanitize_context(company_context or "")
    no_context = not bool(clean_context)

    # Create analysis row
    analysis = RFPAnalysis(
        id=aid,
        user_id=current_user.id,
        status="processing",
        original_name=original_name,
        file_type=ext,
        company_context=clean_context or None,
        no_context=no_context,
    )
    db.add(analysis)
    await db.commit()

    # Parse document synchronously (fast — in-process, no LLM)
    try:
        parsed = parse_rfp_document(file_bytes, ext)
    except ValueError as exc:
        await db.execute(
            sa_delete(RFPAnalysis).where(RFPAnalysis.id == aid)
        )
        await db.commit()
        raise HTTPException(status_code=422, detail=str(exc))

    # Fire pipeline as background task — uses its own DB session to avoid session boundary issues
    async def _background(doc_text: str):
        async with AsyncSessionLocal() as bg_db:
            await run_rfp_analyzer_pipeline(
                analysis_id=aid,
                document_text=doc_text,
                company_context=clean_context or None,
                no_context=no_context,
                db=bg_db,
                anthropic_client=anthropic_client,
            )

    asyncio.create_task(_background(parsed.text))

    logger.info("rfp_analysis_started", analysis_id=str(aid), user_id=str(current_user.id))
    return StartAnalysisResponse(analysis_id=str(aid), status="processing")


@router.get("/history", response_model=list[AnalysisListItem])
async def list_analyses(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(RFPAnalysis)
        .where(RFPAnalysis.user_id == current_user.id)
        .options(selectinload(RFPAnalysis.requirements))
        .order_by(RFPAnalysis.created_at.desc())
        .limit(50)
    )
    analyses = result.scalars().all()
    return [
        AnalysisListItem(
            analysis_id=str(a.id),
            status=a.status,
            original_name=a.original_name,
            client_name=a.client_name,
            created_at=a.created_at.isoformat(),
            requirement_count=len(a.requirements),
        )
        for a in analyses
    ]


@router.get("/{analysis_id}", response_model=AnalysisReportOut)
async def get_analysis(
    analysis_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    analysis = await _get_owned_analysis(analysis_id, current_user, db)
    return _analysis_to_report(analysis)


@router.patch("/{analysis_id}/scope")
async def override_scope(
    analysis_id: str,
    body: OverrideRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if body.scope not in ("in", "conditional", "out"):
        raise HTTPException(status_code=400, detail="scope must be 'in', 'conditional', or 'out'")

    analysis = await _get_owned_analysis(analysis_id, current_user, db)

    # Find the requirement by req_id string (not DB UUID)
    req_row = next((r for r in analysis.requirements if r.req_id == body.req_id), None)
    if not req_row:
        raise HTTPException(status_code=404, detail=f"Requirement '{body.req_id}' not found in this analysis")

    if req_row.classification is None:
        raise HTTPException(status_code=409, detail="Classification not yet available — pipeline may still be processing")

    # Apply override — this is the source of truth, never recalculated (spec rule 9.4)
    req_row.classification.user_override = body.scope
    req_row.classification.override_reason = body.reason
    await db.commit()

    return {"req_id": body.req_id, "scope": body.scope, "saved": True}


@router.delete("/{analysis_id}")
async def delete_analysis(
    analysis_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # Verify ownership before delete
    await _get_owned_analysis(analysis_id, current_user, db)
    aid = uuid.UUID(analysis_id)
    await db.execute(sa_delete(RFPAnalysis).where(RFPAnalysis.id == aid))
    await db.commit()
    return {"deleted": True}


@router.post("/{analysis_id}/export/pdf")
async def export_pdf(
    analysis_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    analysis = await _get_owned_analysis(analysis_id, current_user, db)
    if analysis.status != "complete":
        raise HTTPException(status_code=409, detail="Analysis is not complete yet")

    pdf_bytes = _render_pdf(analysis)
    filename = f"rfp-analysis-{analysis.original_name.rsplit('.', 1)[0]}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/{analysis_id}/export/docx")
async def export_docx(
    analysis_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    analysis = await _get_owned_analysis(analysis_id, current_user, db)
    if analysis.status != "complete":
        raise HTTPException(status_code=409, detail="Analysis is not complete yet")

    docx_bytes = _render_docx(analysis)
    filename = f"rfp-analysis-{analysis.original_name.rsplit('.', 1)[0]}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Export renderers ──────────────────────────────────────────────────────────

def _render_pdf(analysis: RFPAnalysis) -> bytes:
    """Generate a PDF report using reportlab (already in requirements.txt)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph(f"RFP Analysis — {analysis.original_name}", styles["h1"]))
    story.append(Spacer(1, 0.5*cm))

    # Client profile
    if analysis.client_name:
        story.append(Paragraph("Client Profile", styles["h2"]))
        profile_data = [
            ["Client", analysis.client_name or "—"],
            ["Country", analysis.country or "—"],
            ["Sector", analysis.sector or "—"],
            ["Tender ID", analysis.tender_id or "—"],
            ["Deadline", analysis.submission_deadline or "—"],
        ]
        t = Table(profile_data, colWidths=[4*cm, 12*cm])
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F3F0F9")),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.5*cm))

    # Scope map
    story.append(Paragraph("Scope Classification", styles["h2"]))
    scope_data = [["Req ID", "Requirement", "Scope", "Confidence"]]
    for req in analysis.requirements:
        c = req.classification
        if c:
            effective_scope = c.user_override or c.scope
            conf_str = f"{int((c.confidence or 0) * 100)}%" if c.confidence else "—"
            scope_data.append([req.req_id, req.text[:80], effective_scope.upper(), conf_str])

    if len(scope_data) > 1:
        t2 = Table(scope_data, colWidths=[2*cm, 10*cm, 2.5*cm, 2.5*cm])
        t2.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2D1252")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ]))
        story.append(t2)

    doc.build(story)
    return buf.getvalue()


def _render_docx(analysis: RFPAnalysis) -> bytes:
    """Generate a DOCX report using python-docx (already in requirements.txt)."""
    from docx import Document as DocxDocument

    docx = DocxDocument()

    # Title
    docx.add_heading(f"RFP Analysis — {analysis.original_name}", 0)

    # Client profile
    if analysis.client_name:
        docx.add_heading("Client Profile", 1)
        table = docx.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Field", "Value"
        for field, value in [
            ("Client", analysis.client_name),
            ("Country", analysis.country or "—"),
            ("Sector", analysis.sector or "—"),
            ("Tender ID", analysis.tender_id or "—"),
            ("Deadline", analysis.submission_deadline or "—"),
        ]:
            row = table.add_row().cells
            row[0].text, row[1].text = field, value or "—"

    # Scope table
    docx.add_heading("Scope Classification", 1)
    if analysis.requirements:
        t = docx.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Req ID", "Requirement", "Scope", "Confidence"]):
            hdr[i].text = h
        for req in analysis.requirements:
            c = req.classification
            if c:
                row = t.add_row().cells
                row[0].text = req.req_id
                row[1].text = req.text[:120]
                row[2].text = (c.user_override or c.scope).upper()
                row[3].text = f"{int((c.confidence or 0) * 100)}%" if c.confidence else "—"

    buf = io.BytesIO()
    docx.save(buf)
    return buf.getvalue()
